#native imports 
import os

#third party imports
import pandas as pd
import docx


def read_supplementary_file(file_path:str)-> list[pd.DataFrame]:
    """
    Reads a supplementary file based on its extension.
    Supports CSV, TSV, Excel, DOCX (tables), and TXT files.
    """
    print("Starting to read in the supplementary file")
    final_df_list = []
    extension = os.path.splitext(file_path)[1].lower()
    if extension == '.csv':
        df = pd.read_csv(file_path)
        final_df_list.append(df)
    elif extension == '.tsv':
        df = pd.read_csv(file_path, delimiter='\t')
        final_df_list.append(df)
    elif extension in ['.xls', '.xlsx']:
        df = pd.read_excel(file_path)
        final_df_list.append(df)
    elif extension == '.docx':
        doc = docx.Document(file_path)
        tables = doc.tables
        if tables:
            all_tables = []
        # Iterate through each table in the document
            for table in tables:
                # Create a DataFrame structure with empty strings, sized by the number of rows and columns in the table
                df = [['' for _ in range(len(table.columns))] for _ in range(len(table.rows))]
                # Iterate through each row in the current table
                for i, row in enumerate(table.rows):
                    # Iterate through each cell in the current row
                    for j, cell in enumerate(row.cells):
                        # If the cell has text, store it in the corresponding DataFrame position
                        if cell.text:
                            df[i][j] = cell.text #take first line and put as column names
                # Convert the list of lists (df) to a pandas DataFrame and add it to the tables list
                
                all_tables.append(pd.DataFrame(df))
              
                # Print the list of DataFrames representing the tables
                for df in all_tables:
                #TODO should also add some check to see if the first line is text
                    df.columns = df.iloc[0] #take first line and put as column names
                    final_df_list.append(df)
                
        else:
            with open(file_path.replace('.docx', '.txt'), 'w', encoding='utf-8') as txt_file:
                for para in doc.paragraphs:
                    txt_file.write(para.text + '\n')
            raise ValueError("DOCX file does not contain tables. Saved as TXT for further processing.")
    elif extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as txt_file:
            data = txt_file.readlines()
        df = pd.DataFrame(data, columns=['Content'])
        final_df_list.append(df)
    else:
        raise ValueError(f"Unsupported file extension: {extension}")
    return final_df_list

